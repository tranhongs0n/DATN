from docx import Document
doc = Document(r'D:\DATN\thesis\DocOutput\DATN_Report_OnGoing.docx')
print('Total paragraphs:', len(doc.paragraphs))
for p in doc.paragraphs[-20:]:
    if p.text.strip():
        print(p.text[:100])

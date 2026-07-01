import docx
from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fix_images_toc(doc_path):
    doc_path = Path(doc_path)
    print(f"Loading {doc_path.name}...")
    doc = docx.Document(doc_path)
    
    modified = False
    
    for p in doc.paragraphs:
        drawings = p._p.xpath('.//w:drawing')
        if not drawings:
            continue
            
        text = p.text.strip()
        style_name = p.style.name
        
        if text != '':
            print(f"Tách hình khỏi chữ: '{text[:30]}...'")
            new_p = p.insert_paragraph_before()
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            new_p.style = 'Normal'
            
            new_run = new_p.add_run()
            for drawing in drawings:
                new_run._r.append(drawing)
            
            modified = True
            
        elif text == '' and style_name == 'Hình':
            print("Đổi style dòng chứa ảnh sang Normal")
            p.style = 'Normal'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            modified = True

    if modified:
        try:
            doc.save(doc_path)
            print(f"\nSUCCESS! Đã tách hình thành công: {doc_path.name}")
        except PermissionError:
            alt_path = doc_path.parent / (doc_path.stem + "_fixed_images" + doc_path.suffix)
            doc.save(alt_path)
            print(f"\nWARNING: File đang mở, đã lưu bản sửa lỗi tại: {alt_path}")
    else:
        print("\nKhông tìm thấy lỗi hình ảnh nào cần tách.")

if __name__ == "__main__":
    fix_images_toc(r"D:\DATN\thesis\DocOutput\DATN_Report_V2.docx")

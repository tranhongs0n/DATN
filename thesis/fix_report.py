import docx

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

doc = docx.Document("DocOutput/DATN_Report.docx")

# First, modify the second header to say "Chương 5"
for i in range(326, 330):
    if "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN" in doc.paragraphs[i].text:
        doc.paragraphs[i].text = "Chương 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN"
        doc.paragraphs[i].style = "Heading 1"
        break

# Then delete paragraphs 321 to 325
# To do this safely, we find the exact paragraphs to delete so we don't mess up if indexes shifted
paras_to_delete = []
for i in range(320, 330):
    text = doc.paragraphs[i].text.strip()
    if text == "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN" and i < 326:
        paras_to_delete.append(doc.paragraphs[i])
    elif "Lời cảm ơn" in text or "Em xin chân thành cảm ơn" in text or "Em cũng gửi lời tri ân" in text:
        paras_to_delete.append(doc.paragraphs[i])
        
for p in paras_to_delete:
    delete_paragraph(p)

doc.save("DocOutput/DATN_Report_Fixed.docx")
print("Fixed DATN_Report.docx successfully.")

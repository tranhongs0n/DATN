import sys
from docx import Document
from docx.oxml.ns import qn

def inspect_doc(doc_path):
    doc = Document(doc_path)
    body = doc.element.body
    print(f"--- Inspecting {doc_path} ---")
    counts = {}
    for p in doc.paragraphs:
        s = p.style.name if p.style else "None"
        counts[s] = counts.get(s, 0) + 1
        
    print("Styles used by paragraph count:")
    for k, v in sorted(counts.items(), key=lambda x: -x[1]):
        print(f"  {k}: {v}")
        
    # Pick a few Normal paragraphs to see their run properties
    print("\nSample Normal/Normal (Web) runs:")
    c = 0
    for p in doc.paragraphs:
        if p.style and ("Normal" in p.style.name):
            if c < 3 and p.text.strip():
                print(f"Text: {p.text[:50]}")
                for r in p.runs:
                    rPr = r._element.find(qn('w:rPr'))
                    if rPr is not None:
                        fonts = rPr.find(qn('w:rFonts'))
                        sz = rPr.find(qn('w:sz'))
                        color = rPr.find(qn('w:color'))
                        f_val = fonts.get(qn('w:ascii')) if fonts is not None else 'None'
                        sz_val = sz.get(qn('w:val')) if sz is not None else 'None'
                        c_val = color.get(qn('w:val')) if color is not None else 'None'
                        print(f"  Run: font={f_val}, sz={sz_val}, color={c_val}")
                c += 1

print("ORIGINAL TEMPLATE:")
inspect_doc("Ref/HoangThaiDuy_DATN.docx")
print("\nGENERATED DOC:")
inspect_doc("DocOutput/DATN_Report.docx")

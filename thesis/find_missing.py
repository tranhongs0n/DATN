import docx

doc = docx.Document("DocOutput/DATN_Report.docx")
for i, p in enumerate(doc.paragraphs):
    if "Missing" in p.text or "subtitle" in p.text or "thiếu" in p.text.lower():
        print(f"Para {i}: {p.text}")

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if "Missing" in p.text or "subtitle" in p.text:
                    print(f"Table cell: {p.text}")

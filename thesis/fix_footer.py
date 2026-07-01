import sys
from pathlib import Path
import docx
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    t = create_element('w:t')
    t.text = "1"
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)

def fix_footer(doc_path):
    doc_path = Path(doc_path)
    print(f"Loading {doc_path.name}...")
    doc = docx.Document(doc_path)
    
    # 1. Enable TOC Update
    element = doc.settings.element
    updateFields = element.find(ns.qn('w:updateFields'))
    if updateFields is None:
        updateFields = OxmlElement('w:updateFields')
        element.append(updateFields)
    updateFields.set(ns.qn('w:val'), 'true')
    print("Enabled TOC auto-update.")
    
    # 2. Add page numbers to footers
    for i, section in enumerate(doc.sections):
        footer = section.footer
        
        # Clear existing text safely
        for p in footer.paragraphs:
            p.text = ""
        
        if not footer.paragraphs:
            p = footer.add_paragraph()
        else:
            p = footer.paragraphs[0]
            
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        add_page_number(run)
    print("Added page numbers to footers.")

    # Save
    try:
        doc.save(doc_path)
        print(f"\nSUCCESS! Đã cập nhật Footer & TOC cho file: {doc_path}")
    except PermissionError:
        alt_path = doc_path.parent / (doc_path.stem + "_fixed_footer" + doc_path.suffix)
        doc.save(alt_path)
        print(f"\nWARNING: Không thể ghi đè {doc_path.name} vì file đang mở trong Word.")
        print(f"Đã lưu bản cập nhật sang file tạm: {alt_path}")

if __name__ == "__main__":
    fix_footer(r"D:\DATN\thesis\DocOutput\DATN_Report_V2.docx")

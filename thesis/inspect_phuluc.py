import docx

doc = docx.Document("DocOutput/DATN_Report_PhuLuc.docx")
print("Total paras:", len(doc.paragraphs))
for i in range(len(doc.paragraphs) - 15, len(doc.paragraphs)):
    p = doc.paragraphs[i]
    print(f"Para {i}: [{p.style.name}] {p.text}")

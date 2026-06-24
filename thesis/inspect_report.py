import docx

doc = docx.Document("DocOutput/DATN_Report.docx")
print(f"Total paragraphs: {len(doc.paragraphs)}")

for i in range(len(doc.paragraphs)-50, len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if "KẾT LUẬN VÀ HƯỚNG" in p.text.upper() or "LỜI CẢM ƠN" in p.text.upper() or "CẢM ƠN" in p.text.upper():
        print(f"Para {i}: [{p.style.name}] {p.text}")

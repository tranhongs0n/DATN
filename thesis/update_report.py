import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = docx.Document("DocOutput/DATN_Report.docx")

# 1. Update KẾT QUẢ DỰ KIẾN paragraphs
# Find where "KẾT QUẢ DỰ KIẾN:" is
kq_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "KẾT QUẢ DỰ KIẾN:" in p.text:
        kq_idx = i
        break

if kq_idx != -1:
    # Delete the old list items until an empty line or next section
    idx = kq_idx + 1
    while idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        text = p.text.strip()
        if text.startswith("- Nền tảng") or text.startswith("- Bộ cơ sở") or text.startswith("- Bảng điều khiển"):
            # delete it
            el = p._element
            el.getparent().remove(el)
            p._p = p._element = None
        else:
            break

    # Insert the new ones after kq_idx
    new_results = [
        "- Hoàn thiện quyển đồ án tốt nghiệp với đầy đủ cơ sở lý thuyết về RAG, cơ sở dữ liệu vector và quy trình phân tích thiết kế hệ thống.",
        "- Xây dựng được website tích hợp chatbot trợ lý ảo đáp ứng đầy đủ các yêu cầu đã đề ra.",
        "- Tiến hành triển khai thử nghiệm website tích hợp trợ lý ảo trên môi trường thực tế, kiểm tra tính chính xác của câu trả lời dựa trên tập dữ liệu mẫu.",
        "- Đánh giá chất lượng câu trả lời của hệ thống theo các metric tiêu chuẩn: Context Precision, Context Recall và Answer Relevance đạt kết quả từ 0.8 (80%) trở lên; riêng chỉ số Faithfulness đạt mức tiệm cận 1.0 (100%)."
    ]
    
    # We insert in reverse order to keep them right after kq_idx
    target_p = doc.paragraphs[kq_idx]
    for res in reversed(new_results):
        new_p = docx.text.paragraph.Paragraph(OxmlElement("w:p"), doc)
        new_p.text = res
        new_p.style = target_p.style # keep style
        target_p._element.addnext(new_p._element)

# 2. Update TIẾN ĐỘ THỰC HIỆN table
tien_do_tbl = None
for tbl in doc.tables:
    if len(tbl.rows) > 0:
        text = "".join(tbl.rows[0].cells[0].text)
        if "TT" in text and len(tbl.columns) >= 4:
            # Check headers
            header = "".join(c.text for c in tbl.rows[0].cells)
            if "Nội dung công việc" in header:
                tien_do_tbl = tbl
                break

if tien_do_tbl is not None:
    # clear existing rows except header
    for r in tien_do_tbl.rows[1:]:
        el = r._element
        el.getparent().remove(el)
        r._tr = r._element = None
        
    tien_do_contents = [
        ("1", "01/04/2026 - 15/04/2026", "Nghiên cứu lý thuyết, thu thập, tiền xử lý dữ liệu và phân tích thiết kế kiến trúc hệ thống.", "Cơ sở dữ liệu chuẩn hóa, tài liệu thiết kế hệ thống."),
        ("2", "16/04/2026 - 15/05/2026", "Lập trình xây dựng pipeline RAG, cơ sở dữ liệu vector và phát triển giao diện tích hợp.", "Pipeline RAG hoàn chỉnh"),
        ("3", "16/05/2026 - 15/06/2026", "Xây dựng bộ dữ liệu kiểm thử, đo lường đánh giá hệ thống theo các metric và tinh chỉnh tối ưu.", "Bảng kết quả đánh giá metric, hệ thống được tối ưu."),
        ("4", "16/06/2026 - 30/06/2026", "Phân tích kết quả thực nghiệm, viết báo cáo tổng kết và chuẩn bị tài liệu bảo vệ.", "Báo cáo đồ án hoàn chỉnh.")
    ]
    
    for cols in tien_do_contents:
        row_cells = tien_do_tbl.add_row().cells
        for i, cell_text in enumerate(cols):
            row_cells[i].text = cell_text
            # center first two columns
            if i < 2:
                for p in row_cells[i].paragraphs:
                    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

doc.save("DocOutput/DATN_Report_Updated.docx")
print("Updated KẾT QUẢ DỰ KIẾN and TIẾN ĐỘ THỰC HIỆN successfully.")

import os
import re
import sys
import urllib.request
from pathlib import Path
from PIL import Image

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parent
INPUT_DIR = BASE_DIR / "Document_Planing" / "Content"
TEMPLATE_PATH = BASE_DIR / "Ref" / "HoangThaiDuy_DATN.docx"
OUTPUT_PATH = BASE_DIR / "DocOutput" / "DATN_Report_V2.docx"
IMAGES_DIR = BASE_DIR / "DocOutput" / "Images"

PROJECT_INFO = {
    "ho_va_ten":   "TRẦN HỒNG SƠN",
    "ten_de_tai":  "NGHIÊN CỨU VÀ XÂY DỰNG HỆ THỐNG TRỢ LÝ ẢO HỖ TRỢ TUYỂN SINH CHO TRƯỜNG ĐẠI HỌC THỦY LỢI DỰA TRÊN KỸ THUẬT RAG",
    "nganh":       "Công nghệ thông tin",
    "ma_so":       "7480201",
    "nguoi_hd_1":  "TS. Lý Anh Tuấn",
    "nam":         "2026",
    "ten_chinh":   "Trần Hồng Sơn",
    "mssv":        "1951060985",
    "lop":         "61TH1",
    "sdt":         "0339882230",
    "email":       "1951060985@e.tlu.edu.vn",
    "ngay_bat_dau":"31/03/2026",
    "ngay_ket_thuc":"30/06/2026",
}

CHAPTER_FILES = {
    "1_Chuong1.md",
    "2_Chuong2.md",
    "3_Chuong3.md",
    "4_Chuong4.md",
}

sys.stdout.reconfigure(encoding="utf-8")


# ---------------------------------------------------------------------------
# MERMAID
# ---------------------------------------------------------------------------
def fetch_mermaid_png(mermaid_code: str, counter: int) -> str | None:
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    out_path = IMAGES_DIR / f"mermaid_{counter}.png"
    try:
        url = "https://kroki.io/mermaid/png"
        req = urllib.request.Request(
            url, data=mermaid_code.encode("utf-8"), method="POST"
        )
        req.add_header("Content-Type", "text/plain")
        req.add_header("User-Agent", "Mozilla/5.0")
        with urllib.request.urlopen(req, timeout=30) as resp:
            out_path.write_bytes(resp.read())
        print(f"  -> Mermaid {counter}: OK ({out_path.name})")
        return str(out_path)
    except Exception as e:
        print(f"  -> Mermaid {counter}: FAILED ({e})")
        return None


# ---------------------------------------------------------------------------
# FIX #6 & #7: STYLES - Redefine template styles and strip direct formatting
# ---------------------------------------------------------------------------
def normalize_document_formatting(doc: Document) -> None:
    """Redefine styles to perfectly match DocxFormatRule.md and strip direct overrides."""
    # 1. Redefine styles at the document level
    style_specs = [
        ('Normal', 13, False, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 10, 0, 1.5, False),
        ('Heading 1', 14, True, False, WD_ALIGN_PARAGRAPH.LEFT, 24, 24, 1.0, True),
        ('Heading 2', 13, True, False, WD_ALIGN_PARAGRAPH.LEFT, 6, 12, 1.0, False),
        ('Heading 3', 13, True, True, WD_ALIGN_PARAGRAPH.LEFT, 6, 12, 1.0, False),
        ('Heading 4', 13, False, True, WD_ALIGN_PARAGRAPH.LEFT, 6, 12, 1.0, False),
        ('Caption', 12, False, True, WD_ALIGN_PARAGRAPH.CENTER, 6, 6, 1.0, False),
        ('Hình', 12, False, True, WD_ALIGN_PARAGRAPH.CENTER, 6, 6, 1.0, False),
        ('Bảng', 12, False, True, WD_ALIGN_PARAGRAPH.CENTER, 6, 6, 1.0, False),
        ('List Paragraph', 13, False, False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5, False),
        ('Bullet', 13, False, False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5, False)
    ]
    
    for style_name, size, bold, italic, align, space_before, space_after, line_spacing, all_caps in style_specs:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            if hasattr(style, 'font'):
                style.font.name = 'Times New Roman'
                style.font.size = Pt(size)
                style.font.bold = bold
                style.font.italic = italic
                style.font.color.rgb = RGBColor(0, 0, 0)
                if all_caps:
                    style.font.all_caps = True
            if hasattr(style, 'paragraph_format'):
                style.paragraph_format.alignment = align
                style.paragraph_format.space_before = Pt(space_before)
                style.paragraph_format.space_after = Pt(space_after)
                style.paragraph_format.line_spacing = line_spacing
                
    # 2. Strip direct formatting from all runs so they inherit the clean styles
    target_styles = {s[0] for s in style_specs}
    for p in doc.paragraphs:
        if p.style and p.style.name in target_styles:
            for r in p.runs:
                rPr = r._element.find(qn('w:rPr'))
                if rPr is not None:
                    # Remove manual overrides
                    for tag in ['w:color', 'w:rFonts', 'w:sz', 'w:szCs', 'w:b', 'w:bCs', 'w:i', 'w:iCs']:
                        for el in rPr.findall(qn(tag)):
                            rPr.remove(el)

    print("  -> Normalize document formatting: OK (Redefined styles and stripped direct formatting)")


# ---------------------------------------------------------------------------
# NUMBERING PATCH
# ---------------------------------------------------------------------------
def patch_numbering_trailing_dots(doc: Document) -> None:
    """Sua lvlText cua abstractNumId=20 de them dau cham cuoi cho H2/H3/H4."""
    try:
        numbering = doc.part.numbering_part.numbering_definitions._numbering
    except Exception:
        return

    num_to_abstract: dict[str, str] = {}
    for child in numbering:
        tag = child.tag.split("}")[1]
        if tag == "num":
            nid = child.get(qn("w:numId"))
            abs_el = child.find(qn("w:abstractNumId"))
            if abs_el is not None:
                num_to_abstract[nid] = abs_el.get(qn("w:val"))

    target_abs = num_to_abstract.get("28")
    if not target_abs:
        return

    for child in numbering:
        tag = child.tag.split("}")[1]
        if tag == "abstractNum" and child.get(qn("w:abstractNumId")) == target_abs:
            for lvl in child.findall(qn("w:lvl")):
                ilvl = lvl.get(qn("w:ilvl"))
                if ilvl == "0":
                    continue
                lvlText_el = lvl.find(qn("w:lvlText"))
                if lvlText_el is not None:
                    val = lvlText_el.get(qn("w:val"), "")
                    if val and not val.endswith("."):
                        lvlText_el.set(qn("w:val"), val + ".")
            print("  -> Patch numbering: OK (trailing dots added)")
            break


# ---------------------------------------------------------------------------
# TEXT REPLACEMENT HELPERS
# ---------------------------------------------------------------------------
def replace_paragraph_text(para, old: str, new: str) -> None:
    full = para.text
    if old not in full:
        return
    new_full = full.replace(old, new)
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_full
    else:
        para.add_run(new_full)


def fill_table_cell(cell, text: str) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


# ---------------------------------------------------------------------------
# FIX #1 + #2: COMPREHENSIVE FRONT MATTER PATCHING
# ---------------------------------------------------------------------------
def patch_front_matter(doc: Document) -> None:
    """Replace ALL personal info across the entire front matter.
    Covers: cover pages, NHIEM VU, DE CUONG, LỜI CAM ĐOAN sections."""
    placeholders = [
        # Cover page 1
        ("HOÀNG THÁI DUY", PROJECT_INFO["ho_va_ten"].upper()),
        # Cover page 2
        ("Hoàng Thái Duy", PROJECT_INFO["ten_chinh"]),
        # MSSV
        ("1951060657", PROJECT_INFO["mssv"]),
        # Class
        ("61TH2", PROJECT_INFO["lop"]),
        # Title variants (handle line-split)
        ("XÂY DỰNG HỆ THỐNG TRỢ LÝ ẢO HƯỚNG NGHIỆP CHO SINH VIÊN NGÀNH CNTT ĐẠI HỌC THỦY LỢI DỰA TRÊN RAG",
         PROJECT_INFO["ten_de_tai"]),
        ("Xây dựng hệ thống trợ lý ảo hướng nghiệp cho sinh viên ngành CNTT Đại học Thủy Lợi dựa trên RAG",
         "Nghiên cứu và xây dựng hệ thống trợ lý ảo hỗ trợ tuyển sinh cho Trường Đại học Thủy Lợi dựa trên kỹ thuật RAG"),
        # Title on multi-line cover
        ("XÂY DỰNG HỆ THỐNG TRỢ LÝ ẢO HƯỚNG NGHIỆP", "NGHIÊN CỨU VÀ XÂY DỰNG HỆ THỐNG TRỢ LÝ ẢO"),
        ("CHO SINH VIÊN NGÀNH CNTT ĐẠI HỌC THỦY LỢI", "HỖ TRỢ TUYỂN SINH CHO TRƯỜNG ĐẠI HỌC THỦY LỢI"),
        ("DỰA TRÊN RAG", "DỰA TRÊN KỸ THUẬT RAG"),
        # Instructor
        ("TS. Nguyễn Văn Nam", PROJECT_INFO["nguoi_hd_1"]),
        # Year
        ("2025", PROJECT_INFO["nam"]),
        # Email
        ("1951060657@e.tlu.edu.vn", PROJECT_INFO["email"]),
        # Phone
        ("0339882230", PROJECT_INFO["sdt"]),  # already correct in template
        # Date range (ĐỀ CƯƠNG section)
        ("31/03/2025", PROJECT_INFO["ngay_bat_dau"]),
        ("30/06/2025", PROJECT_INFO["ngay_ket_thuc"]),
    ]
    for para in doc.paragraphs:
        for old, new in placeholders:
            if old in para.text:
                replace_paragraph_text(para, old, new)

    # Also patch inside tables (the NHIỆM VỤ table cells)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in placeholders:
                        if old in para.text:
                            replace_paragraph_text(para, old, new)

    print("  -> Patch front matter: OK")


# ---------------------------------------------------------------------------
# REMOVE OLD CONTENT (from "PHẦN MỞ ĐẦU" onward)
# ---------------------------------------------------------------------------
def remove_guide_content(doc: Document) -> None:
    """Xoa toan bo noi dung huong dan tu body element 'PHẦN MỞ ĐẦU' tro di."""
    body = doc.element.body
    start_delete_idx = -1
    for idx, child in enumerate(body):
        tag = child.tag.split("}")[1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            if p.text.strip() == "PHẦN MỞ ĐẦU":
                start_delete_idx = idx
                break

    if start_delete_idx != -1:
        to_delete = list(body)[start_delete_idx:]
        for el in to_delete:
            body.remove(el)
        print(f"  -> Removed {len(to_delete)} old content elements")


# ---------------------------------------------------------------------------
# FLOATING SHAPES
# ---------------------------------------------------------------------------
def remove_floating_shapes(doc: Document) -> None:
    body = doc.element.body
    anchors = body.xpath('//*[local-name()="anchor"]')
    for anchor in anchors:
        drawing = anchor.getparent()
        if drawing is not None:
            run = drawing.getparent()
            if run is not None:
                run.remove(drawing)
    print(f"  -> Removed {len(anchors)} floating shapes")


# ---------------------------------------------------------------------------
# LỜI CẢM ƠN
# ---------------------------------------------------------------------------
def replace_loi_cam_on(doc: Document, md_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    cam_on_match = re.search(
        r"## Lời cảm ơn\s*\n(.*?)(?=\n---|\n#|\Z)", text, re.DOTALL
    )
    if not cam_on_match:
        return

    cam_on_text = cam_on_match.group(1).strip()
    body = doc.element.body

    loi_cam_on_idx = None
    for idx, child in enumerate(body):
        tag = child.tag.split("}")[1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            if "LỜI CÁM ƠN" in p.text or "LỜI CẢM ƠN" in p.text:
                loi_cam_on_idx = idx
                break

    if loi_cam_on_idx is None:
        return

    children_list = list(body)
    placeholder_els = []
    for i in range(loi_cam_on_idx + 1, min(loi_cam_on_idx + 10, len(children_list))):
        child = children_list[i]
        tag = child.tag.split("}")[1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            if p.style.name in ("Content", "Normal") and p.text.strip():
                placeholder_els.append(child)
        else:
            break

    for el in placeholder_els:
        body.remove(el)

    anchor = list(body)[loi_cam_on_idx]
    for line in cam_on_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        new_p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), "Content")
        pPr.append(pStyle)
        new_p.append(pPr)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = stripped
        r.append(t)
        new_p.append(r)
        anchor.addnext(new_p)
        anchor = new_p

    # Also replace "Hoàng Thái Duy" in the signature line
    for idx2, child in enumerate(body):
        tag = child.tag.split("}")[1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            if "Hoàng Thái Duy" in p.text:
                replace_paragraph_text(p, "Hoàng Thái Duy", PROJECT_INFO["ten_chinh"])

    print("  -> Loi cam on: OK")


# ---------------------------------------------------------------------------
# FIX #2: TÓM TẮT ĐỀ TÀI - Correct placement
# ---------------------------------------------------------------------------
def replace_tom_tat(doc: Document, md_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")

    body = doc.element.body
    # Find "TÓM TẮT ĐỀ TÀI" heading
    start_idx = -1
    for idx, child in enumerate(body):
        tag = child.tag.split("}")[1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            if p.text.strip() == "TÓM TẮT ĐỀ TÀI":
                start_idx = idx
                break

    if start_idx == -1:
        print("  -> Tom tat: SKIPPED (heading not found)")
        return

    # Find next major section boundary. 
    end_idx = -1
    children_list = list(body)
    for idx in range(start_idx + 1, len(children_list)):
        child = children_list[idx]
        tag = child.tag.split("}")[1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            t = p.text.strip()
            # Stop at any of these known section landmarks
            if t in ("TIẾN ĐỘ THỰC HIỆN", "KẾT QUẢ DỰ KIẾN",
                     "GÁY BÌA ĐỒ ÁN TỐT NGHIỆP, KHÓA LUẬN TỐT NGHIỆP",
                     "LỜI CAM ĐOAN", "LỜI CÁM ƠN", "LỜI CẢM ƠN",
                     "PHẦN MỞ ĐẦU"):
                end_idx = idx
                break
            # Also stop at Heading 1 style
            if p.style and p.style.name == "Heading 1":
                end_idx = idx
                break
        elif tag == "sectPr":
            end_idx = idx
            break

    if end_idx == -1:
        end_idx = len(children_list)

    # Remove old content between heading and boundary
    to_remove = children_list[start_idx + 1:end_idx]
    for el in to_remove:
        body.remove(el)
    removed_count = len(to_remove)

    # Insert new content after the TÓM TẮT heading
    anchor = list(body)[start_idx]
    lines = text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Skip markdown heading marker
        if stripped.startswith("# "):
            continue

        new_p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), "Content")
        pPr.append(pStyle)
        # Add justify
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "both")
        pPr.append(jc)
        new_p.append(pPr)

        # Handle bold (**...**) markers
        parts = re.split(r"(\*\*.+?\*\*)", stripped)
        for part in parts:
            if not part:
                continue
            r = OxmlElement("w:r")
            if part.startswith("**") and part.endswith("**"):
                rPr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rPr.append(b)
                r.append(rPr)
                t_el = OxmlElement("w:t")
                t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                t_el.text = part[2:-2]
                r.append(t_el)
            else:
                t_el = OxmlElement("w:t")
                t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                t_el.text = part
                r.append(t_el)
            new_p.append(r)

        anchor.addnext(new_p)
        anchor = new_p

    print(f"  -> Tom tat de tai: OK (removed {removed_count} old elements)")


# ---------------------------------------------------------------------------
# FIX #1: NHIỆM VỤ ĐỒ ÁN - Replace ALL sections comprehensively
# ---------------------------------------------------------------------------
def replace_nhiem_vu(doc: Document, md_path: Path) -> None:
    body = doc.element.body
    children_list = list(body)
    
    # 3. Patch References in Nhiem Vu
    # Find "2- CÁC TÀI LIỆU"
    tai_lieu_idx = -1
    for idx, child in enumerate(children_list):
        if child.tag.endswith("p") and "CÁC TÀI LIỆU" in "".join(child.itertext()):
            tai_lieu_idx = idx
            break
            
    if tai_lieu_idx != -1:
        # Delete old list paragraphs until we hit "3 - NỘI DUNG"
        idx = tai_lieu_idx + 1
        while idx < len(children_list):
            if children_list[idx].tag.endswith("p") and "NỘI DUNG CÁC PHẦN" in "".join(children_list[idx].itertext()):
                break
            body.remove(children_list[idx])
            children_list.pop(idx)
            
        # Insert actual references
        refs = [
            "[1] P. Lewis et al., \"Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks,\" NeurIPS, 2020.",
            "[2] A. Vaswani et al., \"Attention Is All You Need,\" NeurIPS, 2017.",
            "[3] Y. Gao et al., \"Retrieval-Augmented Generation for Large Language Models: A Survey,\" arXiv preprint, 2024.",
            "[4] J. Devlin et al., \"BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding,\" NAACL, 2019.",
            "[5] T. Brown et al., \"Language Models are Few-Shot Learners,\" NeurIPS, 2020.",
            "[6] LangChain AI, \"LangChain Documentation,\" 2024.",
            "[7] Chroma, \"ChromaDB Documentation,\" 2024.",
            "[8] Google, \"Google Gemini API Reference,\" 2024."
        ]
        
        for ref_text in reversed(refs):
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = ref_text
            r.append(t)
            p.append(r)
            
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), 'ListParagraph')
            
            # Numbering properties
            numPr = OxmlElement('w:numPr')
            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), '3')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            numPr.append(ilvl)
            numPr.append(numId)
            
            pPr.append(pStyle)
            pPr.append(numPr)
            p.insert(0, pPr)
            children_list[tai_lieu_idx].addnext(p)

    # Find "3 - NỘI DUNG CÁC PHẦN THUYẾT MINH"
    nd_idx = -1
    for idx, child in enumerate(children_list):
        if child.tag.endswith("p") and "NỘI DUNG CÁC PHẦN" in "".join(child.itertext()):
            nd_idx = idx
            break
            
    if nd_idx != -1:
        table_idx = -1
        for idx in range(nd_idx + 1, min(nd_idx + 5, len(children_list))):
            if children_list[idx].tag.endswith("tbl"):
                table_idx = idx
                break
                
        if table_idx != -1:
            tbl = children_list[table_idx]
            # Keep table, clear existing rows except header, add new rows
            rows = tbl.findall(qn("w:tr"))
            if len(rows) > 1:
                for r in rows[1:]:
                    tbl.remove(r)
            
            # Add new rows from KẾT QUẢ DỰ KIẾN
            contents = [
                ("Tìm hiểu tổng quan về bài toán hướng nghiệp, các công nghệ liên quan (RAG, LLM, Vector Database).", "20%"),
                ("Phân tích thiết kế hệ thống, xác định yêu cầu chức năng, phi chức năng và kiến trúc giải pháp.", "30%"),
                ("Thu thập, tiền xử lý và xây dựng cơ sở dữ liệu hướng nghiệp đa phương thức.", "20%"),
                ("Cài đặt và triển khai hệ thống (chatbot Zalo, Web Dashboard, backend RAG).", "20%"),
                ("Đánh giá kết quả, viết báo cáo tổng kết và chuẩn bị tài liệu thuyết minh.", "10%")
            ]
            
            for content, tyle in contents:
                tr = OxmlElement('w:tr')
                for cell_text in [content, tyle]:
                    tc = OxmlElement('w:tc')
                    p = OxmlElement('w:p')
                    r = OxmlElement('w:r')
                    t = OxmlElement('w:t')
                    t.text = cell_text
                    r.append(t)
                    p.append(r)
                    
                    # Formatting
                    pPr = OxmlElement('w:pPr')
                    jc = OxmlElement('w:jc')
                    jc.set(qn('w:val'), 'center' if '%' in cell_text else 'left')
                    pPr.append(jc)
                    p.insert(0, pPr)
                    
                    tc.append(p)
                    tr.append(tc)
                tbl.append(tr)

    # 5. Giáo viên hướng dẫn từng phần
    gv_idx = -1
    for idx, child in enumerate(children_list):
        if child.tag.endswith("p") and "GIÁO VIÊN HƯỚNG DẪN" in "".join(child.itertext()):
            gv_idx = idx
            break
            
    if gv_idx != -1:
        table_idx = -1
        for idx in range(gv_idx + 1, min(gv_idx + 5, len(children_list))):
            if children_list[idx].tag.endswith("tbl"):
                table_idx = idx
                break
                
        if table_idx != -1:
            tbl = children_list[table_idx]
            rows = tbl.findall(qn("w:tr"))
            if len(rows) > 1:
                for r in rows[1:]:
                    tbl.remove(r)
            
            # Add one row for the whole project
            tr = OxmlElement('w:tr')
            for cell_text in ["Toàn bộ dự án", "TS. Lý Anh Tuấn"]:
                tc = OxmlElement('w:tc')
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = cell_text
                r.append(t)
                p.append(r)
                
                # Formatting
                pPr = OxmlElement('w:pPr')
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center')
                pPr.append(jc)
                p.insert(0, pPr)
                
                tc.append(p)
                tr.append(tc)
            tbl.append(tr)

    # 6. TIẾN ĐỘ THỰC HIỆN
    tien_do_tbl = None
    for tbl in doc.element.body.findall(qn("w:tbl")):
        rows = tbl.findall(qn("w:tr"))
        if rows:
            text = "".join(rows[0].itertext())
            if "Nội dung công việc" in text:
                tien_do_tbl = tbl
                break

    if tien_do_tbl is not None:
        rows = tien_do_tbl.findall(qn("w:tr"))
        if len(rows) > 1:
            for r in rows[1:]:
                tien_do_tbl.remove(r)
        
        tien_do_contents = [
            ("1", "01/04/2026 - 15/04/2026", "Nghiên cứu lý thuyết, thu thập, tiền xử lý dữ liệu và phân tích thiết kế kiến trúc hệ thống.", "Cơ sở dữ liệu chuẩn hóa, tài liệu thiết kế hệ thống."),
            ("2", "16/04/2026 - 15/05/2026", "Lập trình xây dựng pipeline RAG, cơ sở dữ liệu vector và phát triển giao diện tích hợp.", "Pipeline RAG hoàn chỉnh"),
            ("3", "16/05/2026 - 15/06/2026", "Xây dựng bộ dữ liệu kiểm thử, đo lường đánh giá hệ thống theo các metric và tinh chỉnh tối ưu.", "Bảng kết quả đánh giá metric, hệ thống được tối ưu."),
            ("4", "16/06/2026 - 30/06/2026", "Phân tích kết quả thực nghiệm, viết báo cáo tổng kết và chuẩn bị tài liệu bảo vệ.", "Báo cáo đồ án hoàn chỉnh.")
        ]
        
        for cols in tien_do_contents:
            tr = OxmlElement('w:tr')
            for i_col, cell_text in enumerate(cols):
                tc = OxmlElement('w:tc')
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = cell_text
                r.append(t)
                p.append(r)
                
                pPr = OxmlElement('w:pPr')
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center' if i_col < 2 else 'left')
                pPr.append(jc)
                p.insert(0, pPr)
                
                tc.append(p)
                tr.append(tc)
            tien_do_tbl.append(tr)

    print("  -> Nhiem vu do an: OK")


# ---------------------------------------------------------------------------
# FIX #5: DANH MỤC VIẾT TẮT - With explicit table borders
# ---------------------------------------------------------------------------
def _make_tbl_borders() -> OxmlElement:
    """Create a tblBorders element with all borders set to single black."""
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    return tblBorders


def replace_danh_muc_viet_tat(doc: Document, md_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    
    # Parse sections and their tables
    sections = re.split(r"^## \d+\.\s+(.+)$", text, flags=re.MULTILINE)
    
    body = doc.element.body
    dmvt_idx = None
    for idx, child in enumerate(body):
        tag = child.tag.split("}")[1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            if "DANH MỤC CÁC TỪ VIẾT TẮT" in p.text:
                dmvt_idx = idx
                break

    if dmvt_idx is None:
        return

    # Remove old placeholder content (paragraphs and tables after the heading)
    children_list = list(body)
    placeholder_els = []
    for i in range(dmvt_idx + 1, min(dmvt_idx + 30, len(children_list))):
        child = children_list[i]
        tag = child.tag.split("}")[1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            if p.style.name in ("Content", "Normal") and p.text.strip():
                placeholder_els.append(child)
            elif not p.text.strip():
                placeholder_els.append(child)
        elif tag == "tbl":
            placeholder_els.append(child)
        else:
            break

    for el in placeholder_els:
        body.remove(el)

    # Parse all markdown tables
    tables_md = re.findall(r"((?:\|.+\|[ \t]*\n)+)", text)
    if not tables_md:
        return

    anchor = list(body)[dmvt_idx]

    # Insert section headers and tables
    section_titles = []
    for i in range(1, len(sections), 2):
        section_titles.append(sections[i])

    table_idx = 0
    for sec_idx, title in enumerate(section_titles):
        # Add section sub-heading
        sub_p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), "Normal")
        pPr.append(pStyle)
        # Add spacing
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "240")
        spacing.set(qn("w:after"), "120")
        pPr.append(spacing)
        sub_p.append(pPr)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr.append(b)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = f"{sec_idx + 1}. {title}"
        r.append(t)
        sub_p.append(r)
        anchor.addnext(sub_p)
        anchor = sub_p

        # Add the corresponding table
        if table_idx < len(tables_md):
            table_md_str = tables_md[table_idx]
            table_idx += 1

            rows_raw = [r.strip() for r in table_md_str.strip().splitlines()]
            rows_data = []
            for row_str in rows_raw:
                if re.match(r"^\|[-:\s|]+\|$", row_str):
                    continue
                cells = [re.sub(r"\*\*(.+?)\*\*", r"\1", c.strip())
                         for c in row_str.strip("|").split("|")]
                rows_data.append(cells)
            if not rows_data:
                continue
            ncols = max(len(r) for r in rows_data)

            tbl = OxmlElement("w:tbl")
            tblPr = OxmlElement("w:tblPr")
            tblStyle = OxmlElement("w:tblStyle")
            tblStyle.set(qn("w:val"), "TableGrid")
            tblW = OxmlElement("w:tblW")
            tblW.set(qn("w:w"), "0")
            tblW.set(qn("w:type"), "auto")
            tblPr.append(tblStyle)
            tblPr.append(tblW)
            # FIX #5: Explicit borders
            tblPr.append(_make_tbl_borders())
            tbl.append(tblPr)

            for r_idx, row_cells in enumerate(rows_data):
                tr = OxmlElement("w:tr")
                for c_idx in range(ncols):
                    cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
                    tc = OxmlElement("w:tc")
                    
                    # Add cell borders too
                    tcPr = OxmlElement("w:tcPr")
                    tcBorders = OxmlElement("w:tcBorders")
                    for edge in ("top", "left", "bottom", "right"):
                        border = OxmlElement(f"w:{edge}")
                        border.set(qn("w:val"), "single")
                        border.set(qn("w:sz"), "4")
                        border.set(qn("w:space"), "0")
                        border.set(qn("w:color"), "000000")
                        tcBorders.append(border)
                    tcPr.append(tcBorders)
                    tc.append(tcPr)
                    
                    p_el = OxmlElement("w:p")
                    pPr_el = OxmlElement("w:pPr")
                    pStyle_el = OxmlElement("w:pStyle")
                    pStyle_el.set(qn("w:val"), "Normal")
                    pPr_el.append(pStyle_el)
                    p_el.append(pPr_el)
                    r_el = OxmlElement("w:r")
                    rPr_el = OxmlElement("w:rPr")
                    # Set font
                    rFonts = OxmlElement("w:rFonts")
                    rFonts.set(qn("w:ascii"), "Times New Roman")
                    rFonts.set(qn("w:hAnsi"), "Times New Roman")
                    rPr_el.append(rFonts)
                    sz_el = OxmlElement("w:sz")
                    sz_el.set(qn("w:val"), "26")  # 13pt
                    rPr_el.append(sz_el)
                    if r_idx == 0:
                        b_el = OxmlElement("w:b")
                        rPr_el.append(b_el)
                    r_el.append(rPr_el)
                    t_el = OxmlElement("w:t")
                    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    t_el.text = cell_text
                    r_el.append(t_el)
                    p_el.append(r_el)
                    tc.append(p_el)
                    tr.append(tc)
                tbl.append(tr)

            anchor.addnext(tbl)
            anchor = tbl

            # Add spacing paragraph after table
            sep_p = OxmlElement("w:p")
            anchor.addnext(sep_p)
            anchor = sep_p

    print("  -> Danh muc viet tat: OK (with borders)")


# ---------------------------------------------------------------------------
# FIX #3 + #4: Set updateFields so Word refreshes TOC/TOF on open
# ---------------------------------------------------------------------------
def enable_update_fields(doc: Document) -> None:
    """Set w:updateFields to true in document settings,
    so Word refreshes all TOC/TOF fields when opening."""
    settings_el = doc.settings.element
    # Check if updateFields already exists
    existing = settings_el.find(qn("w:updateFields"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
    else:
        uf = OxmlElement("w:updateFields")
        uf.set(qn("w:val"), "true")
        settings_el.append(uf)
    print("  -> Enable updateFields: OK (TOC/TOF will refresh on open)")


# ---------------------------------------------------------------------------
# INLINE FORMATTING
# ---------------------------------------------------------------------------
def parse_inline(text: str) -> list[tuple[str, bool, bool, bool]]:
    """Phan tich inline formatting. Tra ve list (text, bold, italic, code)."""
    result = []
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|`([^`]+)`)"
    )
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            result.append((text[last:m.start()], False, False, False))
        raw = m.group(1)
        if raw.startswith("***"):
            result.append((m.group(2), True, True, False))
        elif raw.startswith("**"):
            result.append((m.group(3), True, False, False))
        elif raw.startswith("*"):
            result.append((m.group(4), False, True, False))
        else:
            result.append((m.group(5), False, False, True))
        last = m.end()
    if last < len(text):
        result.append((text[last:], False, False, False))
    return result


def add_formatted_paragraph(
    doc: Document,
    text: str,
    style_name: str,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> docx.text.paragraph.Paragraph:
    para = doc.add_paragraph(style=style_name)
    if align is not None:
        para.alignment = align
    for seg_text, bold, italic, code in parse_inline(text):
        if not seg_text:
            continue
        run = para.add_run(seg_text)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
    return para


def add_caption(doc: Document, text: str, style_name: str = "Caption"):
    try:
        para = doc.add_paragraph(style=style_name)
    except KeyError:
        para = doc.add_paragraph(style="Caption")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True


def add_word_table(doc: Document, rows_data: list[list[str]]) -> None:
    ncols = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=ncols)
    tbl.style = "Table Grid"
    for r_idx, row_cells_data in enumerate(rows_data):
        for c_idx in range(ncols):
            cell_text = row_cells_data[c_idx] if c_idx < len(row_cells_data) else ""
            cell = tbl.rows[r_idx].cells[c_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for seg_text, bold, italic, is_code in parse_inline(cell_text):
                run = p.add_run(seg_text)
                run.bold = bold or (r_idx == 0)
                run.italic = italic
                if is_code:
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
    tbl_pr = tbl._tbl.tblPr
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tbl_pr.append(jc)


def strip_md_numbering(text: str, level: int) -> str:
    # Return text as-is to preserve manual numbering
    return text


# ---------------------------------------------------------------------------
# MAIN CONTENT PROCESSOR
# ---------------------------------------------------------------------------
def preprocess_table_captions(lines: list[str]) -> list[str]:
    new_lines = []
    i = 0
    skip_indices = set()
    while i < len(lines):
        if i in skip_indices:
            i += 1
            continue
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("|"):
            table_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j])
                j += 1
            
            caption_idx = -1
            k = j
            while k < len(lines):
                k_stripped = lines[k].strip()
                if not k_stripped:
                    k += 1
                    continue
                if re.match(r"^\*Bảng\s+[\d\w\.-]+[\.:].+\*\s*$", k_stripped):
                    caption_idx = k
                break
            
            if caption_idx != -1:
                new_lines.append(lines[caption_idx])
                new_lines.append("")
                new_lines.extend(table_lines)
                skip_indices.add(caption_idx)
                i = j
                continue
        
        new_lines.append(line)
        i += 1
    return new_lines


# ---------------------------------------------------------------------------
# MAIN CONTENT PROCESSOR
# ---------------------------------------------------------------------------
def process_content_file(
    doc: Document,
    md_path: Path,
    is_numbered: bool,
    mermaid_counter: list[int],
) -> None:
    raw = md_path.read_text(encoding="utf-8")
    lines = preprocess_table_captions(raw.splitlines())

    in_mermaid = False
    mermaid_code = ""
    in_code = False
    code_lines: list[str] = []
    pending_table_caption: str | None = None
    table_buffer: list[str] = []
    in_table = False

    def flush_table():
        nonlocal pending_table_caption, in_table, table_buffer
        rows_raw = [r.strip() for r in table_buffer]
        rows_data = []
        for r in rows_raw:
            if re.match(r"^\|[-:\s|]+\|$", r):
                continue
            cells = [
                re.sub(r"\*\*(.+?)\*\*", r"\1", c.strip())
                for c in r.strip("|").split("|")
            ]
            rows_data.append(cells)
        if rows_data:
            if pending_table_caption:
                add_caption(doc, pending_table_caption, style_name="Bảng")
                pending_table_caption = None
            add_word_table(doc, rows_data)
        table_buffer.clear()
        in_table = False

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if in_mermaid:
            if stripped.startswith("```"):
                in_mermaid = False
                png_path = fetch_mermaid_png(mermaid_code, mermaid_counter[0])
                mermaid_counter[0] += 1
                if png_path and os.path.exists(png_path):
                    from docx.shared import Inches
                    
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Check dimensions to avoid exceeding half page height
                    with Image.open(png_path) as img:
                        w, h = img.size
                        aspect = w / h
                    
                    # Max width: 5.5 inches, max height: ~4.5 inches (half page)
                    if aspect < (5.5 / 4.5):
                        # Taller than wide, constrain by height
                        p.add_run().add_picture(png_path, height=Inches(4.5))
                    else:
                        # Wider than tall, constrain by width
                        p.add_run().add_picture(png_path, width=Inches(5.5))
            else:
                mermaid_code += line + "\n"
            i += 1
            continue

        if in_code:
            if stripped.startswith("```"):
                in_code = False
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run("\n".join(code_lines))
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                code_lines.clear()
            else:
                code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("```mermaid"):
            if in_table:
                flush_table()
            in_mermaid = True
            mermaid_code = ""
            i += 1
            continue

        if stripped.startswith("```"):
            if in_table:
                flush_table()
            in_code = True
            code_lines.clear()
            i += 1
            continue

        if stripped.startswith("|"):
            in_table = True
            table_buffer.append(stripped)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("> "):
            content_parts = [stripped[2:]]
            while i + 1 < len(lines) and lines[i + 1].strip().startswith("> "):
                i += 1
                content_parts.append(lines[i].strip()[2:])
            p = add_formatted_paragraph(
                doc, " ".join(content_parts), "Normal", WD_ALIGN_PARAGRAPH.JUSTIFY
            )
            p.paragraph_format.left_indent = Pt(24)
            i += 1
            continue

        if re.match(r"^\*Bảng\s+[\d\w\.-]+[\.:].+\*\s*$", stripped):
            pending_table_caption = stripped.strip("*").strip()
            i += 1
            continue

        if re.match(r"^\*(Hình|Bảng)\s+[\d\w\.-]+[\.:].+\*\s*$", stripped):
            caption_text = stripped.strip("*").strip()
            style_name = "Hình" if caption_text.startswith("Hình") else "Bảng"
            add_caption(doc, caption_text, style_name=style_name)
            i += 1
            continue

        if stripped == "---":
            doc.add_page_break()
            i += 1
            continue

        level = 0
        heading_text = ""
        if stripped.startswith("#### "):
            level, heading_text = 4, stripped[5:]
        elif stripped.startswith("### "):
            level, heading_text = 3, stripped[4:]
        elif stripped.startswith("## "):
            level, heading_text = 2, stripped[3:]
        elif stripped.startswith("# "):
            level, heading_text = 1, stripped[2:]

        if level > 0:
            heading_text = strip_md_numbering(heading_text, level)
            heading_text = re.sub(r"\*\*(.+?)\*\*", r"\1", heading_text).strip()

            if is_numbered:
                style_name = f"Heading {level}"
                try:
                    doc.add_paragraph(heading_text, style=style_name)
                except KeyError:
                    p = doc.add_paragraph(heading_text)
                    p.runs[0].bold = True
            else:
                if level == 1:
                    # Use Heading 1 for non-numbered chapters too
                    # (they will get proper formatting from fix_heading_styles)
                    try:
                        doc.add_paragraph(heading_text, style="Heading 1")
                    except KeyError:
                        p = doc.add_paragraph(heading_text)
                        p.runs[0].bold = True
                else:
                    p = add_formatted_paragraph(doc, heading_text, "Normal")
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(13)
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:]
            try:
                add_formatted_paragraph(doc, bullet_text, "Bullet")
            except KeyError:
                add_formatted_paragraph(doc, "• " + bullet_text, "List Paragraph")
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            bullet_text = re.sub(r"^\d+\.\s+", "", stripped)
            add_formatted_paragraph(doc, bullet_text, "List Paragraph")
            i += 1
            continue

        if re.match(r"^\[\d+\]\s+", stripped):
            add_formatted_paragraph(doc, stripped, "Normal", WD_ALIGN_PARAGRAPH.JUSTIFY)
            i += 1
            continue

        add_formatted_paragraph(doc, stripped, "Normal", WD_ALIGN_PARAGRAPH.JUSTIFY)
        i += 1

    if in_table:
        flush_table()


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------
def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    if not TEMPLATE_PATH.exists():
        print(f"Loi: Khong tim thay template tai {TEMPLATE_PATH}")
        return

    print(f"Doc template: {TEMPLATE_PATH.name}")
    doc = Document(TEMPLATE_PATH)

    # FIX: Normalize document styles to match template
    print("Normalize document formatting...")
    normalize_document_formatting(doc)

    print("Sua danh so de muc...")
    patch_numbering_trailing_dots(doc)

    # FIX #1: Comprehensive front matter patching
    print("Dien thong tin trang bia...")
    patch_front_matter(doc)

    print("Xoa noi dung huong dan mau...")
    remove_guide_content(doc)

    print("Xoa floating shapes...")
    remove_floating_shapes(doc)

    phan_cuoi_path = INPUT_DIR / "5_PhanCuoi.md"
    if phan_cuoi_path.exists():
        print("Dien Loi cam on...")
        replace_loi_cam_on(doc, phan_cuoi_path)

    # FIX #5: Abbreviation table with borders
    dmvt_path = INPUT_DIR / "00_DanhMucTuVietTat.md"
    if dmvt_path.exists():
        print("Dien Danh muc tu viet tat...")
        replace_danh_muc_viet_tat(doc, dmvt_path)

    # FIX #2: TÓM TẮT - correct placement
    tom_tat_path = INPUT_DIR / "00_TomTat.md"
    if tom_tat_path.exists():
        print("Dien Tom tat do an...")
        replace_tom_tat(doc, tom_tat_path)

    # FIX #1: NHIỆM VỤ - comprehensive replacement
    nhiem_vu_path = INPUT_DIR / "00_NhiemVu.md"
    if nhiem_vu_path.exists():
        print("Dien Nhiem vu do an...")
        replace_nhiem_vu(doc, nhiem_vu_path)

    # FIX #3 + #4: Enable TOC/TOF update on open
    print("Enable TOC/TOF update...")
    enable_update_fields(doc)

    md_files = sorted(INPUT_DIR.glob("*.md"))
    mermaid_counter = [1]

    for md_file in md_files:
        if md_file.name in ["00_DanhMucTuVietTat.md", "00_TomTat.md", "00_NhiemVu.md"]:
            continue
        is_numbered = md_file.name in CHAPTER_FILES
        print(f"Xu ly: {md_file.name}  (numbered={is_numbered})")
        process_content_file(doc, md_file, is_numbered, mermaid_counter)
        doc.add_page_break()

    try:
        doc.save(OUTPUT_PATH)
        print(f"\nHoan tat! File Word: {OUTPUT_PATH}")
    except PermissionError:
        alt_path = OUTPUT_PATH.parent / "DATN_Report_temp.docx"
        doc.save(alt_path)
        print(f"\nWARNING: Khong the ghi de {OUTPUT_PATH.name} vi file dang mo trong Word.")
        print(f"Da luu file tam thoi tai: {alt_path}")


if __name__ == "__main__":
    main()

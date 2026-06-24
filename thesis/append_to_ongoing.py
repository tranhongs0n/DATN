import sys
from pathlib import Path
from docx import Document
from build_doc import process_content_file, CHAPTER_FILES

BASE_DIR = Path(r"D:\DATN\thesis")
INPUT_DIR = BASE_DIR / "Document_Planing" / "Content"
ONGOING_DOC = BASE_DIR / "DocOutput" / "DATN_Report_OnGoing.docx"
OUTPUT_DOC = BASE_DIR / "DocOutput" / "DATN_Report_Appended.docx"

def append_to_ongoing():
    print(f"Loading {ONGOING_DOC.name}...")
    doc = Document(ONGOING_DOC)
    
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run("--- PHẦN NỘI DUNG MỚI ĐƯỢC CHÈN VÀO (TỪ CHƯƠNG 1 ĐẾN CHƯƠNG 5) ---")
    r.bold = True
    doc.add_page_break()
    
    md_files = sorted(INPUT_DIR.glob("*.md"))
    mermaid_counter = [1]
    
    chapters_to_append = ["1_Chuong1.md", "2_Chuong2.md", "3_Chuong3.md", "4_Chuong4.md", "5_PhanCuoi.md"]
    
    for md_file in md_files:
        if md_file.name not in chapters_to_append:
            continue
            
        is_numbered = md_file.name in CHAPTER_FILES
        print(f"Appending {md_file.name} (numbered={is_numbered})...")
        process_content_file(doc, md_file, is_numbered, mermaid_counter)
        doc.add_page_break()
        
    doc.save(OUTPUT_DOC)
    print(f"\nSUCCESS: Mọi dữ liệu đã được append thành công vào cuối file.")
    print(f"File mới được lưu tại: {OUTPUT_DOC}")

if __name__ == "__main__":
    append_to_ongoing()

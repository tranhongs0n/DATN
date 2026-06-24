import docx

def read_docx(path):
    doc = docx.Document(path)
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text.strip())
            
    print("\n\n--- TABLES ---")
    for t in doc.tables:
        for row in t.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            print(" | ".join(row_data))
        print("---------")

read_docx(r"D:\DATN\thesis\Ref\DeCuongDATN_TranHongSon.docx")

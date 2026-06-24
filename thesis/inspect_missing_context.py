import docx

doc = docx.Document("DocOutput/DATN_Report.docx")
indices = [222, 226, 231, 236]

for idx in indices:
    print(f"\n--- Around {idx} ---")
    start = max(0, idx - 2)
    end = min(len(doc.paragraphs), idx + 3)
    for i in range(start, end):
        print(f"Para {i}: {doc.paragraphs[i].text}")

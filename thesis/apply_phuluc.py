import docx
from docx.enum.text import WD_BREAK

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

doc = docx.Document("DocOutput/DATN_Report_PhuLuc.docx")

# Find paragraphs to delete and rename
paras_to_delete = []
found_first_ket_luan = False

for i in range(len(doc.paragraphs)-50, len(doc.paragraphs)):
    p = doc.paragraphs[i]
    text = p.text.strip().upper()
    
    if text == "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN":
        if not found_first_ket_luan:
            paras_to_delete.append(p)
            found_first_ket_luan = True
        else:
            # This is the second one
            p.text = "Chương 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN"
            p.style = "Heading 1"
    elif "LỜI CẢM ƠN" in text or "EM XIN CHÂN THÀNH CẢM ƠN" in text or "EM CŨNG GỬI LỜI TRI ÂN" in text:
        if found_first_ket_luan: # only delete if it's in the duplicate section at the end
            paras_to_delete.append(p)

for p in paras_to_delete:
    delete_paragraph(p)

# Also delete blank paragraphs between the first Ket luan and second ket luan if any
# We know they were around index 319-324. The above loop captures the specific text ones.

# Now append Phụ lục at the very end
p_break = doc.add_paragraph()
p_break.add_run().add_break(WD_BREAK.PAGE)

p_h1 = doc.add_paragraph("PHỤ LỤC")
p_h1.style = "Heading 1"

p_h2_1 = doc.add_paragraph("Phụ lục 1: Hình ảnh giao diện Zalo Bot")
p_h2_1.style = "Heading 2"

p_place_1 = doc.add_paragraph("[Chèn hình ảnh giao diện Zalo Bot tại đây]")
p_place_1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

p_h2_2 = doc.add_paragraph("Phụ lục 2: Hình ảnh giao diện Web Dashboard Admin")
p_h2_2.style = "Heading 2"

p_place_2 = doc.add_paragraph("[Chèn hình ảnh giao diện Web Dashboard Admin tại đây]")
p_place_2.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

doc.save("DocOutput/DATN_Report_PhuLuc.docx")
print("Phu Luc and fixes applied successfully!")

import requests
import os
import time
import logging
from concurrent.futures import ThreadPoolExecutor
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from src.config.settings import settings

logger = logging.getLogger(__name__)

class TLUAdmissionScraper:
    def __init__(self):
        self.base_url = settings.SCRAPER.base_url
        self.list_api = settings.SCRAPER.list_api
        self.detail_api = settings.SCRAPER.detail_api
        self.extensions = tuple(settings.SCRAPER.extensions)
        self.data_dir = settings.DATA_DIR

    def ensure_dir(self, directory):
        if not os.path.exists(directory):
            os.makedirs(directory)

    def get_last_crawl(self):
        last_crawl_file = self.data_dir / "last_crawl.txt"
        if last_crawl_file.exists():
            return last_crawl_file.read_text().strip()
        return ""

    def update_last_crawl(self):
        last_crawl_file = self.data_dir / "last_crawl.txt"
        from datetime import datetime
        last_crawl_file.write_text(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    def check_new(self):
        total_new = 0
        for cat in settings.SCRAPER.categories:
            name = cat.name
            module_id = cat.module_id
            target_dir = self.data_dir / name
            
            data = self.get_articles(1, module_id)
            if not data or not data.get("ListData"):
                continue
            
            for article in data["ListData"]:
                article_id = article.get("Id")
                filename = f"article_{article_id}.docx"
                filepath = os.path.join(target_dir, filename)
                if not os.path.exists(filepath):
                    total_new += 1
        return total_new

    def get_articles(self, page_index, module_id):
        params = {
            "PortalId": 4,
            "ModuleId": module_id,
            "GetType": 0,
            "IsEmagazine": "false",
            "PageIndex": page_index,
            "txtKeyword": ""
        }
        try:
            response = requests.get(self.list_api, params=params, timeout=10)
            response.raise_for_status()
            return response.json()
        except Exception as e:
            logger.error(f"Error fetching page {page_index}: {e}")
            return None

    def get_article_detail(self, article_id):
        params = {"id": article_id}
        try:
            response = requests.get(self.detail_api, params=params, timeout=10)
            response.raise_for_status()
            return response.json()
        except Exception as e:
            logger.error(f"Error fetching article {article_id}: {e}")
            return None

    def download_file(self, url, folder):
        if not url:
            return None
        
        clean_url = url.split('?')[0]
        filename = os.path.basename(clean_url)
        filepath = os.path.join(folder, filename)
        
        if os.path.exists(filepath):
            return "ALREADY_EXISTS"

        logger.info(f"Downloading: {filename}...")
        try:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
            }
            response = requests.get(url, stream=True, timeout=30, headers=headers)
            response.raise_for_status()
            with open(filepath, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            return "SUCCESS"
        except Exception as e:
            logger.error(f"Failed to download {url}: {e}")
            return "FAILED"

    def scrape_category(self, cat_info, limit=None):
        name = cat_info["name"]
        module_id = cat_info["module_id"]
        target_dir = self.data_dir / name
        self.ensure_dir(target_dir)
        
        logger.info(f"Starting scrape for category: {name}")
        
        page_index = 1
        processed_articles = set()
        
        with ThreadPoolExecutor(max_workers=5) as executor:
            while True:
                data = self.get_articles(page_index, module_id)
                if not data or not data.get("ListData"):
                    break
                
                articles = data["ListData"]
                total_count = data.get("TotalCount", 0)
                
                for article in articles:
                    if limit and len(processed_articles) >= limit:
                        break
                    
                    article_id = article.get("Id")
                    if article_id in processed_articles:
                        continue
                    
                    processed_articles.add(article_id)
                    article_name = article.get('Name')
                    logger.info(f"  [{name}] Processing Article: {article_name}")
                    
                    detail = self.get_article_detail(article_id)
                    if not detail or not detail.get("Content"):
                        continue
                        
                    content = detail["Content"]
                    self.save_article_text(article_name, content, target_dir, article_id)
                    
                    soup = BeautifulSoup(content, 'html.parser')
                    links = soup.find_all('a')
                    
                    download_urls = []
                    for link in links:
                        href = link.get('href')
                        if not href:
                            continue
                            
                        clean_href = href.lower().split('?')[0]
                        if clean_href.endswith(self.extensions):
                            file_url = urljoin(self.base_url, href)
                            download_urls.append(file_url)
                    
                    if download_urls:
                        executor.map(lambda url: self.download_file(url, target_dir), download_urls)
                    
                    time.sleep(0.5)

                if len(processed_articles) >= total_count or (limit and len(processed_articles) >= limit):
                    break
                page_index += 1
        
        logger.info(f"Finished scrape for category: {name}. Processed {len(processed_articles)} articles.")

    def save_article_text(self, title, html_content, folder, article_id):
        from docx import Document
        soup = BeautifulSoup(html_content, 'html.parser')
        
        text = soup.get_text('\n', strip=True)
        if not text.strip():
            return False
            
        filename = f"article_{article_id}.docx"
        filepath = os.path.join(folder, filename)
        if os.path.exists(filepath):
            return "ALREADY_EXISTS"
            
        logger.info(f"Saving article text: {filename}...")
        try:
            doc = Document()
            doc.add_heading(title, level=1)
            
            for p in text.split('\n'):
                p_text = p.strip()
                if p_text:
                    doc.add_paragraph(p_text)
                    
            doc.save(filepath)
            return "SUCCESS"
        except Exception as e:
            logger.error(f"Failed to save article text for {article_id}: {e}")
            return "FAILED"

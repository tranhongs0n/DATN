import os
import logging
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, DirectoryLoader
from src.config.settings import settings
from src.core.vector_db import VectorDBManager
from src.utils.document_loader import DocumentLoader

logger = logging.getLogger(__name__)

class IndexingService:
    """Service to handle document loading and indexing workflows."""
    
    def __init__(self, db_manager: VectorDBManager, loader: DocumentLoader):
        self.db_manager = db_manager
        self.loader = loader

    def get_file_status_list(self):
        """Returns a list of files with their indexing status."""
        files = self.loader.get_available_files()
        indexed_files = self.db_manager.get_indexed_files()
        
        file_data = []
        choices = []
        for f in files:
            basename = os.path.basename(f)
            status = "Indexed" if basename in indexed_files else "Pending"
            file_data.append([basename, status, f])
            choices.append(basename)
            
        local_basenames = {os.path.basename(f) for f in files}
        for idx_file in indexed_files:
            if idx_file not in local_basenames:
                file_data.append([idx_file, "Missing Local", "N/A"])
                
        return file_data, choices

    def load_all_from_disk(self):
        """Loads all supported documents from the data directory."""
        files = self.loader.get_available_files()
        logger.info(f"Loading {len(files)} documents for indexing...")
        return self.load_files_by_path(files)

    def load_files_by_path(self, file_paths):
        """Loads specific documents given their file paths."""
        from src.utils.document_loader import SUPPORTED_DOC_EXTENSIONS
        docs = []
        for f_path in file_paths:
            if not f_path.lower().endswith(SUPPORTED_DOC_EXTENSIONS):
                continue
                
            try:
                if f_path.lower().endswith(".pdf"):
                    docs.extend(PyPDFLoader(f_path).load())
                elif f_path.lower().endswith(".docx"):
                    docs.extend(Docx2txtLoader(f_path).load())
            except Exception as e:
                logger.error(f"Error loading file {f_path}: {e}")
        return docs

    def get_full_path_map(self):
        """Returns a mapping of basenames to full file paths."""
        all_files = self.loader.get_available_files(include_assets=True)
        return {os.path.basename(f): f for f in all_files}

    def convert_unsupported_file(self, filename: str, engine: 'MultimodalEngine'):
        """Converts an unsupported file to .docx using Gemini's multimodal power."""
        path_map = self.get_full_path_map()
        src_path = path_map.get(filename)
        
        if not src_path or not os.path.exists(src_path):
            raise FileNotFoundError(f"File {filename} not found.")

        logger.info(f"Converting {filename} to .docx via Gemini...")
        
        prompt = (
            "You are a professional document digitizer. "
            "Please read the attached file (it could be an image, a legacy document, or a presentation) "
            "and extract all its text content. "
            "Preserve the structure (headings, lists, tables) using Markdown formatting. "
            "Return ONLY the extracted text content."
        )
        
        try:
            content = engine.query([src_path], prompt)
            
            from docx import Document
            doc = Document()
            doc.add_heading(f"Converted Content: {filename}", 0)
            
            for line in content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.strip():
                    doc.add_paragraph(line)
            
            dest_basename = os.path.splitext(filename)[0] + "_converted.docx"
            dest_path = settings.DATA_DIR / dest_basename
            doc.save(dest_path)
            
            logger.info(f"Successfully converted {filename} to {dest_basename}")
            return str(dest_basename)
        except Exception as e:
            logger.error(f"Conversion failed for {filename}: {e}")
            raise e
        finally:
            engine.cleanup()
